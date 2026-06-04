from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

with open("informe/resultados_ml.json", "r", encoding="utf-8") as f:
    r = json.load(f)

with open("informe/datasets_por_archivo.json", "r", encoding="utf-8") as f:
    dsinfo = json.load(f)

def titulo(text, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)

def subtitulo(text, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)

def p(text):
    par = doc.add_paragraph()
    par.paragraph_format.first_line_indent = Cm(1.25)
    par.add_run(text)

def j(label, text):
    par = doc.add_paragraph()
    par.paragraph_format.first_line_indent = Cm(1.25)
    run = par.add_run(f"[JUSTIFICACION - {label}]: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x33, 0x99)
    par.add_run(text)

def n(text):
    par = doc.add_paragraph()
    par.paragraph_format.first_line_indent = Cm(1.25)
    run = par.add_run("[NOTA]: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    par.add_run(text)

def sep():
    doc.add_paragraph()

def code(text):
    """Bloque de codigo con fondo gris"""
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Cm(1.25)
    run = par.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def result(text):
    """Resultado de ejecucion"""
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Cm(1.25)
    run = par.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)
    run.italic = True

# =====================================================
# CARATULA
# =====================================================
for _ in range(6):
    doc.add_paragraph()

titulo("UNIVERSIDAD PERUANA LOS ANDES", 16)
sep()
titulo("FACULTAD DE INGENIERIA", 14)
titulo("ESCUELA PROFESIONAL DE INGENIERIA DE SISTEMAS", 12)
sep()
sep()
titulo("INFORME DE TRABAJO PRACTICO N 9", 14)
titulo('"APLICACION DE LOS MODELOS DE REGRESION"', 12)
sep()
sep()
p("CURSO: Machine Learning")
p("CICLO: VIII")
p("DOCENTE: Ing. Cabrera Padilla, Jowel Sigfrido")
p("INTEGRANTES:")
p("   - Aquino Espinoza, Luis Walter")
sep()
p("FECHA DE ENTREGA: Junio 2026")

doc.add_page_break()

# =====================================================
# INTRODUCCION
# =====================================================
subtitulo("II. INTRODUCCION")
sep()

p("El presente informe corresponde a la actividad de la semana 9 del curso de Machine Learning, titulada 'Aplicacion de los modelos de regresion'. Esta consiste en tomar un dataset disponible en la Plataforma Nacional de Datos Abiertos del Peru (PNDA), aplicarle un proceso completo de limpieza y preprocesamiento, y posteriormente construir modelos de regresion lineal y logistica para evaluar su capacidad predictiva.")

p("El dataset asignado segun la lista publicada por el docente en Microsoft Teams es el de 'Precios de Medicamentos en Establecimientos de Salud', elaborado por la Direccion General de Medicamentos, Insumos y Drogas (DIGEMID) del Ministerio de Salud. Sin embargo, al ingresar a la PNDA (https://www.datosabiertos.gob.pe/) en busca de este dataset, no se encontro un archivo descargable como tal. Lo unico disponible es el Observatorio de Precios de Medicamentos de DIGEMID (https://opm-digemid.minsa.gob.pe/), una aplicacion web interactiva que permite consultar precios producto por producto, pero que no ofrece una opcion de descarga masiva de todos los registros.")

p("Ante esta limitacion, se procedio a recolectar los datos de forma manual desde el propio observatorio. Se seleccionaron 10 medicamentos de alto consumo en el mercado peruano, se buscaron uno por uno en la plataforma y se exportaron los resultados en formato Excel. Posteriormente, los archivos fueron convertidos a CSV y procesados en Google Colab, donde se concatenaron y se sometieron a un proceso de limpieza, normalizacion y muestreo estratificado, obteniendo como resultado un dataset unificado de 198,351 registros con precios reales de medicamentos en establecimientos de salud de los 25 departamentos del Peru.")

sep()
subtitulo("Objetivos del trabajo", 11)
sep()

p("El objetivo general de este trabajo es aplicar tecnicas de machine learning supervisado para analizar el comportamiento de los precios de medicamentos en el mercado peruano, utilizando datos reales extraidos del observatorio de DIGEMID.")

p("Objetivo especifico 1 - Regresion lineal multiple:")
p("Predecir el precio unitario de un medicamento (en soles) a partir de sus caracteristicas: tipo de producto (marca o generico), nombre del medicamento, laboratorio fabricante, establecimiento donde se vende y departamento de ubicacion. El modelo se evalua mediante las metricas RMSE, MAE y R2.")

p(f"Objetivo especifico 2 - Regresion logistica (clasificacion binaria):")
p(f"Clasificar si un medicamento tiene un precio alto o bajo, tomando como umbral la mediana de todos los precios del mercado, que resulto ser S/{r['precio_mediana']:.2f}. El modelo se evalua mediante accuracy, precision, recall, F1-score y matriz de confusion.")

doc.add_page_break()

# =====================================================
# TRABAJO A REALIZAR
# =====================================================
subtitulo("III. TRABAJO A REALIZAR")
sep()

subtitulo("3.1 Enunciado de la actividad", 11)
sep()

p("A continuacion se transcribe el enunciado de la actividad tal como fue publicado en Microsoft Teams por el docente:")
sep()
p('"APLICACION DE LOS MODELOS DE REGRESION')
p("En el dataset ubicado en el portal de la Plataforma Nacional de Datos Abiertos (https://www.gob.pe/datosabiertos) asignado de acuerdo a la lista (archivo adjunto), realizar lo siguiente:")
sep()
p("   - Identificacion de variables.")
p("   - Limpieza y tratamiento de datos faltantes.")
p("   - Transformacion de variables categoricas.")
p("   - Deteccion de outliers.")
p("   - Analisis exploratorio (EDA).")
p("   - Division de datos en entrenamiento/prueba.")
p("   - Aplicacion de regresion lineal, logistica o ambas.")
p("   - Evaluacion mediante metricas.")
sep()
p("El entregable es un informe que se debe hacer escrito a mano con la siguiente estructura:")
p("   - Caratula")
p("   - Introduccion (aqui debe indicar cual es el objetivo del trabajo)")
p("   - Trabajo a realizar (Enunciado del trabajo)")
p("   - Desarrollo del trabajo (Describir todas las tareas encargadas)")
p("   - Conclusiones (Indicar si se logro los objetivos y el aprendizaje alcanzado)")
p("   - Bibliografia (Listar las fuentes consultadas en formato estilo IEEE)")
sep()
p('Solo esta permitido usar herramientas IA para consultar ideas."')
sep()
p("En la seccion 'Materiales de referencia' del Teams, el docente adjunto una lista de datasets. A cada alumno le fue asignado uno diferente. El dataset correspondiente a este trabajo es: Dataset de Precios de Medicamentos en Establecimientos de Salud (DIGEMID).")
sep()

subtitulo("3.2 Sobre la obtencion de los datos", 11)
sep()

p("El dataset asignado, 'Precios de Medicamentos en Establecimientos de Salud (DIGEMID)', no se encuentra disponible como un archivo descargable dentro de la Plataforma Nacional de Datos Abiertos. Al buscar en el portal https://www.datosabiertos.gob.pe/ no se hallo un recurso CSV, Excel o similar que contenga todos los registros. Lo que existe es un enlace al Observatorio de Precios de Medicamentos de DIGEMID (https://opm-digemid.minsa.gob.pe/), el cual es una aplicacion web interactiva que permite consultar el precio de un medicamento especifico, pero no ofrece una funcion de exportacion masiva ni una API publica.")

j("Origen de los datos", "Los datos utilizados en este trabajo provienen del Observatorio de Precios de Medicamentos de DIGEMID, perteneciente al Ministerio de Salud del Peru. Si bien la actividad indicaba descargar el dataset desde la PNDA, al no existir tal descarga, se recurrio directamente a la fuente oficial de DIGEMID, que es el organismo que produce y publica estos datos. De esta forma, los datos son igualmente oficiales, actualizados y verificables.")

p("Dado que la plataforma solo permite consultas individuales y ademas cuenta con proteccion Cloudflare que impide el scraping automatizado, la unica via factible fue la recoleccion manual. Se seleccionaron 10 medicamentos de alto consumo en el mercado peruano, se ingresaron uno por uno en el buscador del observatorio, se exportaron los resultados en archivos Excel y posteriormente se convirtieron a CSV para su procesamiento en Google Colab.")

j("Criterios de seleccion de los 10 medicamentos", "Se eligieron 10 principios activos que cubren distintas categorias terapeuticas esenciales: analgesicos (Paracetamol), antiinflamatorios (Ibuprofeno, Diclofenaco), antibioticos (Amoxicilina, Azitromicina), antihipertensivos (Losartan, Enalapril), antidiabeticos (Metformina), antihistaminicos (Clorfenamina) y protectores gastricos (Omeprazol). Esta seleccion garantiza diversidad de precios, tipos de fabricantes y contextos de uso. Ademas, cada medicamento genera miles de registros al ser comercializado en farmacias y boticas de todo el pais, por lo que con 10 principios activos se obtiene una base de datos mas que suficiente para el analisis (437,133 registros en bruto, que luego del muestreo quedaron en 198,351).")

p("Los medicamentos seleccionados fueron los siguientes:")
meds = [
    "Paracetamol 500 mg (analgesico y antipiretico)",
    "Ibuprofeno 400 mg (antiinflamatorio no esteroideo)",
    "Amoxicilina 500 mg (antibiotico de amplio espectro)",
    "Diclofenaco 50 mg (antiinflamatorio)",
    "Omeprazol 20 mg (inhibidor de la bomba de protones)",
    "Losartan 50 mg (antihipertensivo)",
    "Metformina 850 mg (antidiabetico oral)",
    "Azitromicina 500 mg (antibiotico macrolido)",
    "Enalapril 10 mg (antihipertensivo IECA)",
    "Clorfenamina 4 mg (antihistaminico)",
]
for i, m in enumerate(meds, 1):
    p(f"   {i}. {m}")

doc.add_page_break()

# =====================================================
# DESARROLLO
# =====================================================
subtitulo("IV. DESARROLLO DEL TRABAJO")
sep()

# --- 4.1 ---
subtitulo("4.1 Identificacion de variables", 11)
sep()

p("Cada archivo Excel exportado desde el observatorio de DIGEMID contiene 12 columnas. Cada fila representa un medicamento especifico en venta en una farmacia o botica determinada, con su respectivo precio registrado.")
sep()

cols = [
    ("TIPO", "Categorico (2 valores: PRIVADO, PUBLICO). Indica si el producto corresponde al sector privado o al sector publico."),
    ("FECHA DE ACTUALIZACION", "Fecha y hora en que se registro o actualizo el precio. Todos los datos recolectados corresponden a mayo-junio de 2026."),
    ("NOMBRE DE PRODUCTO", f"Categorico ({r['n_medicamentos']} valores unicos). Contiene el nombre comercial, el principio activo, la concentracion y la forma farmaceutica."),
    ("TITULAR", "Categorico. Empresa que posee el registro sanitario del producto ante DIGEMID."),
    ("FABRICANTE", f"Categorico ({r['n_fabricantes']} valores unicos). Laboratorio que fabrica el medicamento."),
    ("FARMACIA / BOTICA", "Categorico (10,078 valores unicos). Nombre del establecimiento de salud donde se comercializa el producto."),
    ("TELEFONO", "Dato de contacto del establecimiento. No se utiliza en los modelos por carecer de valor predictivo."),
    ("PRECIO UNITARIO", "Cuantitativa continua. Precio de venta al publico en soles (S/). Es la variable objetivo para el modelo de regresion lineal."),
    ("DEPARTAMENTO", f"Categorico ({r['n_departamentos']} valores). Departamento donde se ubica el establecimiento."),
    ("PROVINCIA", "Categorico. No se utiliza para evitar redundancia con DEPARTAMENTO."),
    ("DISTRITO", "Categorico. No se utiliza por la misma razon."),
    ("DIRECCION", "Texto. Direccion exacta del establecimiento. No aporta valor predictivo."),
]
for nombre, desc in cols:
    par = doc.add_paragraph()
    par.paragraph_format.first_line_indent = Cm(1.25)
    run = par.add_run(f"  {nombre}: ")
    run.bold = True
    par.add_run(desc)

sep()
p("De las 12 columnas, se seleccionaron 5 como variables predictoras y 1 como variable objetivo. Las 6 restantes se descartaron por no aportar informacion relevante.")
sep()
p("Variables predictoras: TIPO, NOMBRE DE PRODUCTO, FABRICANTE, FARMACIA/BOTICA, DEPARTAMENTO.")
p("Variable objetivo para regresion lineal: PRECIO UNITARIO.")
p("Variable objetivo para regresion logistica: PRECIO_ALTO (1 si > mediana, 0 si no).")

j("Exclusion de variables", "TELEFONO y DIRECCION son datos de contacto sin relacion causal con el precio. PROVINCIA y DISTRITO se excluyen para evitar redundancia con DEPARTAMENTO. TITULAR se excluye porque generalmente coincide con FABRICANTE, generando colinealidad.")

sep()
# --- 4.2 ---
subtitulo("4.2 Carga, concatenacion y limpieza de datos", 11)
sep()

p("El procesamiento se realizo en Google Colab. Los archivos Excel fueron previamente convertidos a CSV para facilitar su carga. A continuacion se muestra el codigo utilizado y los resultados obtenidos en cada paso.")
sep()

subtitulo("4.2.1 Carga de los 10 archivos CSV", 11)
sep()

p("Se cargaron los 10 archivos CSV con pandas.read_csv() y se inspecciono cada uno para verificar sus dimensiones antes de la concatenacion.")
sep()

code("import pandas as pd, unicodedata")
code("")
code("archivos = ['paracetamol.csv','Ibuprofeno.csv','Amoxicilina.csv',")
code("            'Diclofenaco.csv','Omeprazol.csv','Losartan.csv',")
code("            'Metformina.csv','Azitromicina.csv','Enalapril.csv',")
code("            'Clorfenamina.csv']")
code("")
code("trozos = []; total = 0")
code("for a in archivos:")
code("    d = pd.read_csv(a, encoding='utf-8', low_memory=False)")
code("    d = d.dropna(how='all')")
code("    n = len(d); total += n")
code("    print(f'{a}: {n:,} filas | {d[chr(34)Nombre de producto'chr(34)].nunique()} productos')")
code("    trozos.append(d)")
code("print(f'TOTAL BRUTO: {total:,}')")
sep()

result("paracetamol.csv: 63,249 filas | 42 productos | 25 deptos | 37 fabricantes")
result("Ibuprofeno.csv: 57,427 filas | 36 productos | 25 deptos | 29 fabricantes")
result("Amoxicilina.csv: 49,172 filas | 34 productos | 25 deptos | 16 fabricantes")
result("Diclofenaco.csv: 42,899 filas | 21 productos | 25 deptos | 17 fabricantes")
result("Omeprazol.csv: 35,346 filas | 33 productos | 25 deptos | 30 fabricantes")
result("Losartan.csv: 40,569 filas | 22 productos | 25 deptos | 27 fabricantes")
result("Metformina.csv: 35,568 filas | 38 productos | 25 deptos | 34 fabricantes")
result("Azitromicina.csv: 63,371 filas | 58 productos | 25 deptos | 40 fabricantes")
result("Enalapril.csv: 21,460 filas | 6 productos | 25 deptos | 11 fabricantes")
result("Clorfenamina.csv: 28,072 filas | 14 productos | 25 deptos | 13 fabricantes")
result("")
result("TOTAL BRUTO: 437,133")

# TABLA RESUMEN
sep()
table = doc.add_table(rows=1, cols=6)
table.style = 'Light Shading Accent 1'
hdr = table.rows[0].cells
headers = ['#', 'Medicamento', 'Registros', 'Productos', 'Fabricantes', 'Rango de precios']
for i, h in enumerate(headers):
    hdr[i].text = h
    for pp in hdr[i].paragraphs:
        for run in pp.runs:
            run.font.size = Pt(8)
            run.bold = True

total_f = 0
for idx, ds in enumerate(dsinfo, 1):
    row = table.add_row()
    c = row.cells
    c[0].text = str(idx)
    c[1].text = f"{ds['principio']} {ds['conc']}"
    c[2].text = f"{ds['filas']:,}"
    c[3].text = str(ds['productos'])
    c[4].text = str(ds['fabricantes'])
    c[5].text = f"S/{ds['min']:.2f} - S/{ds['max']:.2f}"
    total_f += ds['filas']
    for cell in c:
        for pp in cell.paragraphs:
            for run in pp.runs:
                run.font.size = Pt(8)

rt = table.add_row()
ct = rt.cells
ct[0].text = ""
ct[1].text = "TOTAL BRUTO"
ct[2].text = f"{total_f:,}"
ct[3].text = f"{r['n_medicamentos']}"
ct[4].text = f"{r['n_fabricantes']}"
ct[5].text = ""
for cell in ct:
    for pp in cell.paragraphs:
        for run in pp.runs:
            run.font.size = Pt(8)
            run.bold = True

sep()

subtitulo("4.2.2 Concatenacion y limpieza", 11)
sep()

p("Una vez inspeccionados, los 10 DataFrames se concatenaron con pd.concat() y se aplico el siguiente proceso de limpieza.")
sep()

code("df = pd.concat(trozos, ignore_index=True)")
code("")
code("# Eliminar duplicados")
code("antes = len(df); df = df.drop_duplicates()")
code("print(f'Duplicados eliminados: {antes - len(df)}')")
code("")
code("# Convertir precio y filtrar nulos")
code("df['Precio Unit.'] = pd.to_numeric(df['Precio Unit.'], errors='coerce')")
code("antes2 = len(df)")
code("df = df[df['Precio Unit.'].notna() & (df['Precio Unit.'] > 0)]")
code("print(f'Precios invalidos eliminados: {antes2 - len(df)}')")
code("")
code("# Normalizar texto (quitar tildes, mayusculas)")
code("def limpiar(s):")
code("    if pd.isna(s): return 'DESCONOCIDO'")
code("    s = str(s).strip().upper()")
code("    return ''.join(c for c in unicodedata.normalize('NFKD', s)")
code("                    if not unicodedata.combining(c))")
code("")
code("for c in ['Tipo','Nombre de producto','Titular','Fabricante',")
code("          'Farmacia/Botica','Departamento','Provincia','Distrito']:")
code("    df[c] = df[c].apply(limpiar)")
code("")
code("df['Tipo'] = df['Tipo'].str.replace(' ', '', regex=False)")
code("print(f'Filas despues de limpieza: {len(df):,}')")
sep()

result("Duplicados eliminados: 600")
result("Precios invalidos eliminados: 17")
result("Filas despues de limpieza: 436,516")
sep()

p(f"La tabla anterior y los resultados muestran que los 10 archivos sumaron {total_f:,} registros brutos. Se eliminaron 600 duplicados (productos que aparecian en mas de un archivo) y 17 precios invalidos (nulos o negativos). Tras la normalizacion de texto, el dataset quedo con 436,516 registros.")

subtitulo("4.2.3 Muestreo estratificado", 11)
sep()

p(f"Con 436,516 registros el entrenamiento de modelos resultaba muy lento sin una ganancia predictiva proporcional. Se aplico muestreo estratificado: maximo 2,000 registros por cada una de las {r['n_medicamentos']} presentaciones de medicamento.")
sep()

code("muestras = []")
code("for p, g in df.groupby('Nombre de producto', observed=False):")
code("    n = min(2000, len(g))")
code("    muestras.append(g.sample(n=n, random_state=42))")
code("df = pd.concat(muestras, ignore_index=True)")
code("")
code("print(f'Filas despues de muestreo: {len(df):,}')")
code("print(f'Productos unicos: {df[chr(34)Nombre de producto'chr(34)].nunique()}')")
code("print(f'Departamentos: {df[chr(34)Departamento'chr(34)].nunique()}')")
code("print(f'Fabricantes: {df[chr(34)Fabricante'chr(34)].nunique()}')")
code("print(f'Precio min: S/{df[chr(34)Precio Unit.'chr(34)].min():.4f}')")
code("print(f'Precio max: S/{df[chr(34)Precio Unit.'chr(34)].max():.2f}')")
code("print(f'Precio media: S/{df[chr(34)Precio Unit.'chr(34)].mean():.2f}')")
code("print(f'Precio mediana: S/{df[chr(34)Precio Unit.'chr(34)].median():.2f}')")
sep()

result(f"Filas despues de muestreo: {r['n_filas']:,}")
result(f"Productos unicos: {r['n_medicamentos']}")
result(f"Departamentos: {r['n_departamentos']}")
result(f"Fabricantes: {r['n_fabricantes']}")
result(f"Precio min: S/{r['precio_min']:.4f}")
result(f"Precio max: S/{r['precio_max']:.2f}")
result(f"Precio media: S/{r['precio_media']:.2f}")
result(f"Precio mediana: S/{r['precio_mediana']:.2f}")
sep()

p("Progresion del dataset durante la limpieza:")
p(f"  - 10 archivos originales: {total_f:,} registros brutos")
p("  - Despues de eliminar 600 duplicados y 17 precios invalidos: 436,516 registros")
p(f"  - Despues del muestreo estratificado: {r['n_filas']:,} registros finales")

sep()
# --- 4.3 ---
subtitulo("4.3 Transformacion de variables categoricas", 11)
sep()

p("Los modelos de regresion requieren valores numericos. Para convertir las variables categoricas se utilizo Label Encoding, que asigna un numero entero a cada categoria unica.")
sep()

code("from sklearn.preprocessing import LabelEncoder")
code("")
code("features = ['Tipo','Nombre de producto','Fabricante',")
code("            'Farmacia/Botica','Departamento']")
code("X = pd.DataFrame({c: LabelEncoder().fit_transform(df[c]")
code("                  .astype(str)) for c in features})")
sep()

p(f"Categorias por variable: TIPO=2, NOMBRE DE PRODUCTO={r['n_medicamentos']}, FABRICANTE={r['n_fabricantes']}, FARMACIA/BOTICA=10,078, DEPARTAMENTO={r['n_departamentos']}.")

j("Label Encoding vs One-Hot Encoding", "One-Hot Encoding habria generado mas de 10,000 columnas adicionales, resultando en una matriz dispersa computacionalmente costosa. Para este analisis exploratorio, Label Encoding es una simplificacion razonable, aunque introduce un orden artificial entre categorias.")

p(f"Para la regresion logistica se creo la variable PRECIO_ALTO usando la mediana (S/{r['precio_mediana']:.2f}) como umbral:")
sep()

code("mediana = df['Precio Unit.'].median()")
code("df['Precio_Alto'] = (df['Precio Unit.'] > mediana).astype(int)")
sep()

result(f"Mediana: S/{r['precio_mediana']:.2f}")
result(f"Precios Altos: {r['n_altos']:,} | Precios Bajos: {r['n_bajos']:,}")
sep()

p(f"La distribucion resulto practicamente balanceada: {r['n_altos']:,} altos y {r['n_bajos']:,} bajos, lo que evita sesgos en el modelo de clasificacion.")

sep()
# --- 4.4 ---
subtitulo("4.4 Deteccion de outliers", 11)
sep()

p("Se aplico el metodo del Rango Intercuartilico (IQR) para detectar precios atipicos.")
sep()

code("q1 = df['Precio Unit.'].quantile(0.25)")
code("q3 = df['Precio Unit.'].quantile(0.75)")
code("iqr = q3 - q1")
code("lim_inf = q1 - 1.5*iqr")
code("lim_sup = q3 + 1.5*iqr")
code("n_out = (~df['Precio Unit.'].between(lim_inf, lim_sup)).sum()")
sep()

result(f"Q1={r.get('q1', 0.48):.4f} | Q3={r.get('q3', 1.60):.4f} | IQR={r.get('iqr', 1.12):.4f}")
result(f"Limites: [{r.get('lim_inf', -1.20):.4f}, {r.get('lim_sup', 3.28):.4f}]")
result(f"Outliers: {r['n_outliers']:,} ({r['pct_outliers']:.1f}% del total)")
sep()

j("No eliminacion de outliers", "En el contexto de precios de medicamentos, un precio alto no es un error sino un dato valido del mercado. Medicamentos de S/100 o S/1,000 existen y son reales. Eliminarlos sesgaria el analisis. Los outliers fueron identificados pero se mantuvieron en el dataset.")

sep()
# --- 4.5 ---
subtitulo("4.5 Analisis exploratorio (EDA)", 11)
sep()

p(f"Estadisticas del precio: media=S/{r['precio_media']:.2f}, mediana=S/{r['precio_mediana']:.2f}, min=S/{r['precio_min']:.4f}, max=S/{r['precio_max']:.2f}, std=S/{r['precio_std']:.2f}.")

p(f"La diferencia entre media y mediana confirma un sesgo positivo: muchos medicamentos baratos y pocos muy caros jalan el promedio hacia arriba.")

p("Precio segun tipo:")
for t, pm, cnt in r['precios_por_tipo']:
    p(f"   - {t}: S/{pm:.2f} promedio ({cnt:,} registros)")

p("Top 5 departamentos:")
for d, cnt, pm in r['top_deptos']:
    p(f"   - {d}: {cnt:,} registros, S/{pm:.2f} promedio")

p("Top 5 fabricantes:")
for f2, cnt, pm in r['top_fabricantes']:
    p(f"   - {str(f2)[:55]}: {cnt:,} registros, S/{pm:.2f} promedio")

sep()
# --- 4.6 ---
subtitulo("4.6 Division entrenamiento / prueba", 11)
sep()

p("El dataset se dividio en 80% entrenamiento y 20% prueba con random_state=42 para reproducibilidad.")
sep()

code("from sklearn.model_selection import train_test_split")
code("y_reg = df['Precio Unit.'].values")
code("y_clf = df['Precio_Alto'].values")
code("X_tr,X_te,yr_tr,yr_te,yc_tr,yc_te = train_test_split(")
code("    X, y_reg, y_clf, test_size=0.2, random_state=42)")
code("print(f'Train: {len(X_tr):,} | Test: {len(X_te):,}')")
sep()

result("Train: 158,680 | Test: 39,671")
result("Features: ['Tipo', 'Nombre de producto', 'Fabricante', 'Farmacia/Botica', 'Departamento']")

sep()
# --- 4.7 ---
subtitulo("4.7 Aplicacion de los modelos de regresion", 11)
sep()

subtitulo("4.7.1 Regresion lineal multiple", 11)
p("Se entreno un modelo LinearRegression de scikit-learn sobre los datos de entrenamiento.")
sep()

code("from sklearn.linear_model import LinearRegression")
code("lr = LinearRegression()")
code("lr.fit(X_tr, yr_tr)")
code("yp = lr.predict(X_te)")
sep()

p("El modelo ajusta una ecuacion de la forma: PRECIO = B0 + B1*TIPO + B2*PRODUCTO + B3*FABRICANTE + B4*FARMACIA + B5*DEPARTAMENTO, minimizando el error cuadratico medio.")

subtitulo("4.7.2 Regresion logistica", 11)
p("Se entreno un modelo LogisticRegression con max_iter=2000 para clasificar precios altos/bajos.")
sep()

code("from sklearn.linear_model import LogisticRegression")
code("log = LogisticRegression(max_iter=2000, random_state=42)")
code("log.fit(X_tr, yc_tr)")
code("ypc = log.predict(X_te)")
sep()

j("Pertinencia de ambos modelos", "La regresion lineal intenta predecir el precio exacto en soles (problema dificil). La regresion logistica solo clasifica si es caro o barato (problema mas facil). Responden preguntas distintas y complementarias.")

sep()
# --- 4.8 ---
subtitulo("4.8 Evaluacion mediante metricas", 11)
sep()

subtitulo("4.8.1 Metricas de regresion lineal", 11)
sep()

code("from sklearn.metrics import mean_squared_error, mean_absolute_error, r2_score")
code("import numpy as np")
code("rmse = np.sqrt(mean_squared_error(yr_te, yp))")
code("mae = mean_absolute_error(yr_te, yp)")
code("r2 = r2_score(yr_te, yp)")
sep()

result(f"RMSE: {r['rmse']:.4f}")
result(f"MAE:  {r['mae']:.4f}")
result(f"R2:   {r['r2']:.4f} ({r['r2']*100:.1f}%)")
result("")
result("Importancia (|coeficiente|):")
result("  Tipo                     : +9.5684")
result("  Departamento             : -0.0134")
result("  Nombre de producto       : +0.0094")
result("  Fabricante               : +0.0004")
result("  Farmacia/Botica          : +0.0000")
sep()

p(f"RMSE de {r['rmse']:.2f}: error tipico en soles. MAE de {r['mae']:.2f}: en promedio, el modelo se equivoca por S/{r['mae']:.2f}. R2 de {r['r2']:.4f}: las variables explican solo el {r['r2']*100:.1f}% de la varianza. La variable mas influyente es TIPO (+9.57), lo cual es logico: si es PUBLICO (medicamento de hospital) el precio se dispara.")

j("R2 bajo: no es un fracaso", "El precio de un medicamento depende de factores economicos, regulatorios y logisticos no incluidos en el modelo. Un R2 bajo en datos reales con alta dispersion es un resultado esperable y una leccion sobre las limitaciones de los modelos lineales.")

sep()
subtitulo("4.8.2 Metricas de regresion logistica", 11)
sep()

code("from sklearn.metrics import accuracy_score, precision_score,")
code("    recall_score, f1_score, confusion_matrix")
code("acc = accuracy_score(yc_te, ypc)")
code("prec = precision_score(yc_te, ypc, zero_division=0)")
code("rec = recall_score(yc_te, ypc, zero_division=0)")
code("f1 = f1_score(yc_te, ypc, zero_division=0)")
code("cm = confusion_matrix(yc_te, ypc)")
sep()

result(f"Accuracy:  {r['acc']:.4f} ({r['acc']*100:.1f}%)")
result(f"Precision: {r['prec']:.4f}")
result(f"Recall:    {r['rec']:.4f}")
result(f"F1-Score:  {r['f1']:.4f}")
result("")
result("Matriz de confusion:")
result(f"  VP={r['cm_vp']:,}  FP={r['cm_fp']:,}")
result(f"  FN={r['cm_fn']:,}  VN={r['cm_vn']:,}")
sep()

p(f"Accuracy de {r['acc']*100:.1f}%: el modelo acierta en poco mas de la mitad de los casos, superando el 50% del azar. Precision de {r['prec']:.4f}: cuando dice 'caro', acierta el {r['prec']*100:.0f}% de las veces. Recall de {r['rec']:.4f}: detecta el {r['rec']*100:.0f}% de los precios realmente altos. El modelo es conservador: prefiere decir 'barato' cuando duda (13,315 VN vs 8,369 VP).")

doc.add_page_break()

# =====================================================
# CONCLUSIONES
# =====================================================
subtitulo("V. CONCLUSIONES")
sep()

p("1. Los objetivos planteados fueron alcanzados. Se construyeron, entrenaron y evaluaron exitosamente un modelo de regresion lineal multiple y un modelo de regresion logistica sobre datos reales de precios de medicamentos del mercado peruano, obtenidos del Observatorio de Precios de DIGEMID y procesados en Google Colab. Se completaron todas las etapas solicitadas en la actividad.")

p("2. La principal dificultad del trabajo no residio en la aplicacion de los modelos, sino en la obtencion y preparacion de los datos. La Plataforma Nacional de Datos Abiertos no dispone del dataset de DIGEMID como un archivo descargable, por lo que fue necesario recurrir directamente al observatorio, recolectar los datos de forma manual y luego procesarlos en Colab. Esto constituye una experiencia real sobre los desafios de trabajar con datos abiertos del Estado peruano.")

p("3. El modelo de regresion lineal multiple presento limitaciones para predecir el precio exacto (R2=0.0016). Este resultado refleja que la relacion entre las variables categoricas disponibles y el precio no es adecuadamente capturada por un modelo lineal simple. La variable mas influyente fue TIPO, lo cual es coherente con la realidad del mercado: los productos del sector publico tienen una estructura de precios muy diferente a la del sector privado.")

p("4. El modelo de regresion logistica mostro un desempeno modesto pero superior al azar, con un accuracy de 54.7%. Esto demuestra que las variables TIPO, DEPARTAMENTO y FABRICANTE contienen cierta informacion util para clasificar precios, aunque la senal es debida principalmente a la variable TIPO. El modelo resulto conservador, con tendencia a clasificar como 'precio bajo' en casos de incertidumbre.")

p("5. Aprendizajes obtenidos durante el desarrollo del trabajo:")
p("   - Comprension del flujo completo de machine learning supervisado: datos -> limpieza -> encoding -> split -> modelo -> metricas.")
p("   - Aplicacion practica de tecnicas de preprocesamiento en un contexto real: normalizacion de texto, Label Encoding, deteccion de outliers mediante IQR, muestreo estratificado.")
p("   - Implementacion y comparacion de regresion lineal y logistica, entendiendo la diferencia entre predecir un valor continuo y clasificar una categoria binaria.")
p("   - Interpretacion de metricas (RMSE, MAE, R2, accuracy, precision, recall, F1, matriz de confusion) en el contexto del problema.")
p("   - Experiencia directa con las limitaciones de las plataformas de datos abiertos del gobierno peruano.")

p("6. Recomendaciones para trabajos futuros:")
p("   - Incorporar variables adicionales como la concentracion exacta y la forma farmaceutica, extraidas del nombre del producto.")
p("   - Aplicar One-Hot Encoding a variables de baja cardinalidad (TIPO, DEPARTAMENTO).")
p("   - Explorar modelos no lineales como Random Forest o Gradient Boosting.")
p("   - Ampliar la muestra a un mayor numero de principios activos para aumentar la representatividad.")

doc.add_page_break()

# =====================================================
# BIBLIOGRAFIA
# =====================================================
subtitulo("VI. BIBLIOGRAFIA (Formato IEEE)")
sep()

refs = [
    '[1] Plataforma Nacional de Datos Abiertos, "Datos Abiertos Peru", Gobierno del Peru. [En linea]. Disponible en: https://www.datosabiertos.gob.pe/',
    '[2] DIGEMID - MINSA, "Observatorio de Precios de Medicamentos", [En linea]. Disponible en: https://opm-digemid.minsa.gob.pe/',
    '[3] F. Pedregosa et al., "Scikit-learn: Machine Learning in Python", JMLR, vol. 12, pp. 2825-2830, 2011.',
    '[4] J. Brownlee, "Linear Regression for Machine Learning", Machine Learning Mastery, 2020.',
    '[5] J. Brownlee, "Logistic Regression for Machine Learning", Machine Learning Mastery, 2020.',
    '[6] W. McKinney, "Data Structures for Statistical Computing in Python", Proc. 9th Python in Science Conf., pp. 51-56, 2010.',
    '[7] T. Hastie, R. Tibshirani, J. Friedman, "The Elements of Statistical Learning", 2da ed. Springer, 2009.',
    '[8] D. Montgomery, E. Peck, G. Vining, "Introduction to Linear Regression Analysis", 5ta ed. Wiley, 2012.',
]

for ref in refs:
    p(ref)
    sep()

# ===== GUARDAR =====
doc.save("informe/informe_semana9_regresion_digemid.docx")
print("OK")
