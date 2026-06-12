from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

def titulo(text, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)

def subtitulo(text, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)

def p(text):
    par = doc.add_paragraph()
    par.paragraph_format.first_line_indent = Cm(1.25)
    par.add_run(text)

def sep():
    doc.add_paragraph()

def code(text):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Cm(1.25)
    run = par.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def result(text):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Cm(1.25)
    run = par.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)
    run.italic = True

# =====================================================
# CARATULA
# =====================================================
for _ in range(6):
    doc.add_paragraph()

titulo("UNIVERSIDAD PERUANA LOS ANDES", 16)
sep()
titulo("FACULTAD DE INGENIERIA", 14)
titulo("ESCUELA PROFESIONAL DE INGENIERIA DE SISTEMAS", 12)
sep()
sep()
titulo("INFORME DE PRACTICA DE CLUSTERING", 14)
titulo('"APLICACION DE ALGORITMOS DE CLUSTERING AL DATASET WINE"', 12)
sep()
sep()
p("CURSO: Machine Learning")
p("CICLO: VIII")
p("DOCENTE: Ing. Cabrera Padilla, Jowel Sigfrido")
p("INTEGRANTES:")
p("   - Aquino Espinoza, Luis Walter")
sep()
p("FECHA DE ENTREGA: Junio 2026")

doc.add_page_break()

# =====================================================
# INTRODUCCION
# =====================================================
subtitulo("II. INTRODUCCION")
sep()

p("El clustering o agrupamiento es una tecnica de aprendizaje no supervisado que busca dividir un conjunto de datos en grupos (clusters) donde los elementos dentro de un mismo grupo son similares entre si y diferentes a los de otros grupos. A diferencia de los metodos supervisados, en clustering no se dispone de etiquetas previas; el algoritmo descubre la estructura subyacente de los datos por si mismo.")

p("En esta practica se aplican dos de los algoritmos de clustering mas representativos: K-Means (particional) y AGNES (jerarquico aglomerativo), utilizando el dataset Wine disponible en la libreria scikit-learn. Este dataset contiene 178 muestras de vinos italianos con 13 caracteristicas quimicas (alcohol, acido malico, ceniza, magnesio, fenoles, flavonoides, intensidad de color, etc.) clasificadas en 3 tipos de vino.")

p("El objetivo es determinar el numero optimo de clusters mediante los metodos del codo y del coeficiente de silueta, aplicar ambos algoritmos, visualizar los resultados con PCA y comparar el desempeno de cada uno.")

doc.add_page_break()

# =====================================================
# PROBLEMA PLANTEADO
# =====================================================
subtitulo("III. PROBLEMA PLANTEADO")
sep()

p("Se dispone del dataset Wine de scikit-learn con las siguientes caracteristicas:")
p("   - 178 muestras de vinos")
p("   - 13 atributos quimicos continuos")
p("   - 3 clases reales (class_0, class_1, class_2)")
sep()

p("Se solicita:")
p("   1. Determinar el K optimo mediante el metodo del codo y el coeficiente de silueta.")
p("   2. Aplicar el algoritmo K-Means con el K optimo encontrado.")
p("   3. Aplicar el algoritmo AGNES (clustering jerarquico aglomerativo) con enlace Ward.")
p("   4. Visualizar los resultados mediante PCA (reduccion a 2 dimensiones).")
p("   5. Generar el dendrograma correspondiente al clustering jerarquico.")
p("   6. Comparar y analizar los resultados obtenidos por ambos metodos.")

doc.add_page_break()

# =====================================================
# SOLUCION
# =====================================================
subtitulo("IV. SOLUCION")
sep()

p("A continuacion se presenta la secuencia completa de solucion, documentada mediante el codigo Python utilizado y los resultados obtenidos en cada paso. El procesamiento se realizo en Google Colab.")
sep()

# --- Paso 1 ---
subtitulo("4.1 Carga y exploracion del dataset", 11)
sep()

p("Se carga el dataset Wine, se escalan los datos con StandardScaler (requisito indispensable para clustering) y se aplica PCA para reducir a 2 dimensiones y facilitar la visualizacion.")
sep()

code("import numpy as np, matplotlib.pyplot as plt")
code("from sklearn.datasets import load_wine")
code("from sklearn.preprocessing import StandardScaler")
code("from sklearn.decomposition import PCA")
code("")
code("wine = load_wine()")
code("X, y_real = wine.data, wine.target")
code("X_scaled = StandardScaler().fit_transform(X)")
code("X_pca = PCA(n_components=2).fit_transform(X_scaled)")
sep()

result("Dataset Wine: 178 muestras, 13 caracteristicas, 3 clases reales")
result("Varianza PCA 2D: 55.4%")
result("Clases: class_0, class_1, class_2")
sep()

# --- Paso 2 ---
subtitulo("4.2 Metodo del Codo para determinar K optimo", 11)
sep()

p("Se entrena K-Means con K desde 1 hasta 10 y se registra la inercia (WCSS) en cada caso. El codo se identifica donde la reduccion de inercia se estanca.")
sep()

code("from sklearn.cluster import KMeans")
code("")
code("inercias = []")
code("for k in range(1, 11):")
code("    km = KMeans(n_clusters=k, random_state=42, n_init='auto')")
code("    km.fit(X_scaled)")
code("    inercias.append(km.inertia_)")
sep()

result("K= 1  ->  Inercia = 2314.00")
result("K= 2  ->  Inercia = 1661.68")
result("K= 3  ->  Inercia = 1277.93   <-- CODO")
result("K= 4  ->  Inercia = 1211.75")
result("K= 5  ->  Inercia = 1123.16")
result("...")
result("K=10  ->  Inercia =  879.43")
sep()

p("Se observa una disminucion pronunciada hasta K=3, donde la curva forma un codo. A partir de K=4 la reduccion de inercia es mucho menor. Por tanto, el metodo del codo sugiere K=3 como el numero optimo de clusters, lo cual coincide con la cantidad real de clases del dataset.")
sep()

# --- Paso 3 ---
subtitulo("4.3 Coeficiente de Silueta", 11)
sep()

p("Se calcula el coeficiente de silueta para K desde 2 hasta 10. El valor mas alto indica la mejor separacion entre clusters.")
sep()

code("from sklearn.metrics import silhouette_score")
code("")
code("sil_scores = []")
code("for k in range(2, 11):")
code("    labels = KMeans(n_clusters=k, random_state=42,")
code("                    n_init='auto').fit_predict(X_scaled)")
code("    sil_scores.append(silhouette_score(X_scaled, labels))")
sep()

result("K= 2  ->  Silhouette = 0.2650")
result("K= 3  ->  Silhouette = 0.2849   <-- MEJOR")
result("K= 4  ->  Silhouette = 0.2542")
result("K= 5  ->  Silhouette = 0.1836")
result("...")
result("")
result("K optimo segun Silueta: 3")
sep()

p("El coeficiente de silueta confirma que K=3 es el valor optimo, con un score de 0.2849. Ambos metodos (codo y silueta) coinciden en K=3, que ademas corresponde al numero real de tipos de vino en el dataset.")
sep()

# --- Paso 4 ---
subtitulo("4.4 K-Means con K=3", 11)
sep()

p("Se entrena el modelo K-Means definitivo con 3 clusters y se visualizan los resultados comparados con las etiquetas reales.")
sep()

code("kmeans = KMeans(n_clusters=3, random_state=42, n_init='auto')")
code("labels_kmeans = kmeans.fit_predict(X_scaled)")
code("centroids_pca = pca.transform(kmeans.cluster_centers_)")
sep()

result("Inercia:        1277.93")
result("Iteraciones:    7")
result("Clusters:       [65, 51, 62]")
result("")
result("Matriz de confusion (Real vs K-Means):")
result("  [[ 0  0 59]")
result("   [65  3  3]")
result("   [ 0 48  0]]")
sep()

p("La matriz de confusion muestra que K-Means clasifica correctamente la gran mayoria de las muestras. El cluster 0 de K-Means captura 59 vinos de class_2, el cluster 1 captura 65 de class_0 (con solo 3 errores hacia class_1 y 3 hacia class_2), y el cluster 2 captura 48 de class_1. La pureza alcanzada es del 96.6%.")
sep()

# --- Paso 5 ---
subtitulo("4.5 AGNES - Clustering Jerarquico Aglomerativo", 11)
sep()

p("Se aplica el algoritmo AgglomerativeClustering con enlace Ward (minimiza la varianza intra-cluster). Se utiliza K=3.")
sep()

code("from sklearn.cluster import AgglomerativeClustering")
code("")
code("agg = AgglomerativeClustering(n_clusters=3, linkage='ward')")
code("labels_agg = agg.fit_predict(X_scaled)")
sep()

result("Clusters:       [58, 56, 64]")
result("")
result("Matriz de confusion (Real vs AGNES):")
result("  [[ 0  0 59]")
result("   [58  8  5]")
result("   [ 0 48  0]]")
sep()

p("AGNES tambien logra una buena separacion, aunque con ligeramente mas errores que K-Means (8 falsos en class_0 vs class_1, y 5 en class_0 vs class_2). La pureza es del 92.7%.")
sep()

# --- Paso 6 ---
subtitulo("4.6 Dendrograma", 11)
sep()

p("Se genera el dendrograma del clustering jerarquico usando SciPy con el metodo Ward. La linea roja punteada indica el corte a K=3.")
sep()

code("from scipy.cluster.hierarchy import dendrogram, linkage")
code("")
code("linked = linkage(X_scaled, method='ward')")
code("dendrogram(linked, leaf_rotation=90, leaf_font_size=8)")
code("plt.axhline(y=cut_height, color='r', linestyle='--')")
sep()

p("El dendrograma muestra la estructura jerarquica completa: desde las 178 hojas individuales hasta el cluster raiz. La altura de cada fusion representa la distancia (Ward) entre los grupos que se unen. El corte horizontal en K=3 separa claramente tres grandes ramas correspondientes a los tres tipos de vino.")
sep()

# --- Paso 7 ---
subtitulo("4.7 Comparacion final", 11)
sep()

p("Se comparan ambos algoritmos mediante el coeficiente de silueta y la pureza respecto a las etiquetas reales.")
sep()

result("Silhouette K-Means:  0.2849")
result("Silhouette AGNES:    0.2774")
result("")
result("Pureza K-Means:  96.6%")
result("Pureza AGNES:    92.7%")
sep()

p("Ambos algoritmos obtienen resultados muy similares y satisfactorios. K-Means presenta una ligera ventaja tanto en silueta como en pureza. AGNES, si bien es levemente inferior en metricas, tiene la ventaja de proporcionar el dendrograma como herramienta visual para entender la estructura jerarquica de los datos.")

doc.add_page_break()

# =====================================================
# ANALISIS DE RESULTADOS
# =====================================================
subtitulo("V. ANALISIS DE LOS RESULTADOS OBTENIDOS")
sep()

p("El analisis de los resultados permite extraer las siguientes observaciones:")

p("1. DETERMINACION DE K: Tanto el metodo del codo como el coeficiente de silueta coincidieron en K=3 como el numero optimo de clusters. Esto es consistente con el conocimiento previo del dataset, que efectivamente contiene 3 tipos de vino. La inercia mostro una reduccion significativa de K=1 a K=3 (de 2314 a 1278) y luego se estabilizo, formando un codo claramente identificable. El coeficiente de silueta alcanzo su maximo en K=3 con un valor de 0.2849, confirmando que esta es la particion que mejor equilibra cohesion y separacion.")

p("2. DESEMPENO DE K-MEANS: El algoritmo particional K-Means logro una pureza del 96.6%, clasificando correctamente 172 de las 178 muestras. Los errores se concentraron en la clase class_0, donde 3 muestras fueron asignadas a class_1 y 3 a class_2. Esto sugiere que class_0 comparte cierta similitud quimica con las otras dos clases en algunas muestras, lo cual es esperable en productos naturales como el vino. La inercia final de 1277.93 y las solo 7 iteraciones necesarias para converger indican que el algoritmo encontro rapidamente una solucion de buena calidad.")

p("3. DESEMPENO DE AGNES: El algoritmo jerarquico aglomerativo con enlace Ward obtuvo una pureza del 92.7% (165 de 178 aciertos). Cometio mas errores que K-Means en class_0 (8 y 5 muestras mal clasificadas). Sin embargo, AGNES ofrece una ventaja conceptual importante: el dendrograma permite visualizar la estructura jerarquica completa de los datos, mostrando como se van fusionando las muestras desde el nivel individual hasta formar los 3 clusters finales. Esta visualizacion es particularmente util para entender relaciones de similitud entre grupos.")

p("4. COMPARACION ENTRE ALGORITMOS: K-Means resulto ligeramente superior en este caso (silueta 0.2849 vs 0.2774, pureza 96.6% vs 92.7%). Esto es esperable porque K-Means reasigna iterativamente los puntos a los centroides, refinando la solucion, mientras que AGNES toma decisiones de fusion que son irrevocables. Sin embargo, la diferencia es pequena y ambos metodos demuestran ser efectivos para este dataset. La eleccion entre uno y otro dependera del contexto: K-Means es mas rapido y escalable; AGNES proporciona el dendrograma y no requiere especificar K a priori si se usa un criterio de corte.")

p("5. ROL DEL PCA: La reduccion a 2 dimensiones mediante PCA capturo el 55.4% de la varianza total. Si bien no es un porcentaje extremadamente alto (se pierde un 44.6% de informacion), fue suficiente para visualizar una separacion clara entre los clusters en el espacio bidimensional. Esto demuestra la utilidad del PCA como herramienta de visualizacion en problemas de clustering.")

doc.add_page_break()

# =====================================================
# CONCLUSIONES
# =====================================================
subtitulo("VI. CONCLUSIONES")
sep()

p("1. Se aplicaron exitosamente los algoritmos de clustering K-Means y AGNES sobre el dataset Wine, completando todas las etapas solicitadas: determinacion del K optimo, entrenamiento de los modelos, visualizacion con PCA y dendrograma, y comparacion de resultados.")

p("2. El metodo del codo y el coeficiente de silueta confirmaron K=3 como el numero optimo de clusters, lo cual coincide con el numero real de clases del dataset. Esto valida la efectividad de ambos metodos para la seleccion de K en problemas donde se desconoce el numero de grupos a priori.")

p("3. K-Means demostro un desempeno ligeramente superior con una pureza del 96.6% y un coeficiente de silueta de 0.2849. Su rapidez y simplicidad lo convierten en una excelente primera opcion para tareas de clustering.")

p("4. AGNES con enlace Ward obtuvo resultados muy cercanos (92.7% de pureza, silueta 0.2774) y ofrece la ventaja adicional del dendrograma, que permite visualizar la estructura jerarquica de los datos sin necesidad de definir K de antemano.")

p("5. La practica permitio comprender la diferencia fundamental entre clustering particional (K-Means) y jerarquico (AGNES), asi como la importancia del escalado de datos, la reduccion de dimensionalidad con PCA y el uso de metricas como la inercia, el coeficiente de silueta y la pureza para evaluar la calidad de los agrupamientos.")

doc.add_page_break()

# =====================================================
# BIBLIOGRAFIA
# =====================================================
subtitulo("VII. BIBLIOGRAFIA (Formato IEEE)")
sep()

refs = [
    '[1] F. Pedregosa et al., "Scikit-learn: Machine Learning in Python", Journal of Machine Learning Research, vol. 12, pp. 2825-2830, 2011.',
    '[2] P. Virtanen et al., "SciPy 1.0: Fundamental Algorithms for Scientific Computing in Python", Nature Methods, vol. 17, pp. 261-272, 2020.',
    '[3] J. MacQueen, "Some Methods for Classification and Analysis of Multivariate Observations", Proc. 5th Berkeley Symposium on Mathematical Statistics and Probability, pp. 281-297, 1967.',
    '[4] L. Kaufman y P. J. Rousseeuw, Finding Groups in Data: An Introduction to Cluster Analysis, Wiley, 1990.',
    '[5] T. Hastie, R. Tibshirani y J. Friedman, The Elements of Statistical Learning, 2da ed. Springer, 2009.',
    '[6] P. J. Rousseeuw, "Silhouettes: A graphical aid to the interpretation and validation of cluster analysis", Journal of Computational and Applied Mathematics, vol. 20, pp. 53-65, 1987.',
    '[7] J. D. Hunter, "Matplotlib: A 2D Graphics Environment", Computing in Science & Engineering, vol. 9, no. 3, pp. 90-95, 2007.',
]

for ref in refs:
    p(ref)
    sep()

# ===== GUARDAR =====
doc.save("informe/informe_clustering_wine.docx")
print("DOCX generado: informe/informe_clustering_wine.docx")
